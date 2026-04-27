#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def shade_cell(cell, color):
    """Add background color to a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_test_table(doc, title, rows_data):
    """Create a formatted test table"""
    doc.add_heading(title, level=2)

    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, cell in enumerate(header_cells):
        shade_cell(cell, 'd5e8f0')
        cell.text = rows_data[0][i]
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)

    # Data rows
    for row_idx in range(1, len(rows_data)):
        cells = table.rows[row_idx].cells
        for col_idx, cell in enumerate(cells):
            cell.text = rows_data[row_idx][col_idx]
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

# Create document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title page
title = doc.add_heading('GESTESCOLAR', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(26, 115, 232)
title_run.font.size = Pt(36)

subtitle = doc.add_heading('Plano de Testes - MVP', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.color.rgb = RGBColor(26, 115, 232)
subtitle_run.font.size = Pt(28)

doc.add_paragraph()

desc = doc.add_paragraph('Documento de Validação e Testes de Funcionalidades')
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc.runs[0].font.italic = True
desc.runs[0].font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

date_para = doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.size = Pt(11)

# Page break
doc.add_page_break()

# 1. Introduction
doc.add_heading('1. Introdução', level=1)
doc.add_paragraph('Este documento apresenta o plano de testes abrangente para o MVP (Minimum Viable Product) do GestEscolar. O objetivo é validar todas as funcionalidades críticas do sistema antes do lançamento.')
doc.add_paragraph('O teste deve ser executado de forma sistemática, seguindo cada etapa descrita neste documento, garantindo que todas as funcionalidades funcionem corretamente em diferentes cenários.')

# 2. Prerequisites
doc.add_heading('2. Pré-requisitos', level=1)
for item in [
    'Acesso ao sistema em ambiente de testes',
    'Contas de teste para todos os papéis: Super Admin, Administrativo, Financeiro, Professor e Pai',
    'Base de dados de teste populada com escolas, turmas, alunos e boletos',
    'Acesso ao email para verificar recuperação de senha e mensagens',
    'Navegador web atualizado (Chrome, Firefox ou Safari)'
]:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

# 3. Authentication Tests
doc.add_heading('3. Testes de Autenticação', level=1)

create_test_table(doc, '3.1 Login com Credenciais Válidas', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Navegar para a página de login', 'Página de login é exibida'],
    ['2', 'Inserir email de teste válido', 'Email é aceito no campo'],
    ['3', 'Inserir senha correta', 'Senha é aceita no campo'],
    ['4', 'Clicar em \'Entrar\'', 'Usuário é autenticado e redirecionado para dashboard'],
    ['5', 'Verificar se sessão está ativa', 'Token de sessão válido em sessionStorage/localStorage'],
    ['✓', 'Teste Passado', 'Login bem-sucedido para todos os papéis']
])

create_test_table(doc, '3.2 Login com Credenciais Inválidas', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Navegar para a página de login', 'Página de login é exibida'],
    ['2', 'Inserir email inválido ou senha incorreta', 'Campos aceitam a entrada'],
    ['3', 'Clicar em \'Entrar\'', 'Mensagem de erro: \'Email ou senha inválidos\''],
    ['4', 'Verificar se usuário não é autenticado', 'Usuário permanece na página de login'],
    ['✓', 'Teste Passado', 'Erro é exibido claramente']
])

create_test_table(doc, '3.3 Logout', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Fazer login com credenciais válidas', 'Usuário entra no dashboard'],
    ['2', 'Localizar botão de logout', 'Botão está visível'],
    ['3', 'Clicar em logout', 'Sessão é encerrada'],
    ['4', 'Verificar redirecionamento', 'Página inicial ou login é exibida'],
    ['5', 'Acessar dashboard via URL', 'Sistema redireciona para login'],
    ['✓', 'Teste Passado', 'Logout funciona corretamente']
])

# 4. Password Recovery
doc.add_heading('4. Testes de Recuperação de Senha', level=1)

create_test_table(doc, '4.1 Solicitar Recuperação de Senha', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Clicar em \'Esqueceu sua senha?\'', 'Formulário de recuperação é exibido'],
    ['2', 'Inserir email de uma conta existente', 'Email é aceito'],
    ['3', 'Clicar em \'Enviar Link\'', 'Mensagem de sucesso: \'Email enviado\''],
    ['4', 'Verificar caixa de entrada', 'Email com link de recuperação foi recebido'],
    ['✓', 'Teste Passado', 'Link de recuperação é enviado com sucesso']
])

create_test_table(doc, '4.2 Redefinir Senha via Link', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Clicar no link de recuperação', 'Página de redefinição é exibida'],
    ['2', 'Verificar email pré-preenchido', 'Email aparece no formulário'],
    ['3', 'Inserir nova senha válida', 'Senha é aceita, força é indicada'],
    ['4', 'Confirmar a mesma senha', 'Confirmação é aceita'],
    ['5', 'Clicar em \'Salvar Nova Senha\'', 'Mensagem de sucesso aparece'],
    ['6', 'Login com nova senha', 'Acesso é permitido'],
    ['✓', 'Teste Passado', 'Senha é redefinida corretamente']
])

# 5. Dashboards
doc.add_heading('5. Testes de Dashboards', level=1)

create_test_table(doc, '5.1 Dashboard Super Admin', [
    ['Elemento', 'Esperado', 'Status'],
    ['Acesso', 'Super Admin consegue acessar /superadmin-dashboard', ''],
    ['Navbar', 'Menu com: Escolas, Usuários, Configurações', ''],
    ['Métricas', 'Total de escolas, usuários, transações', ''],
    ['Tabela', 'Lista todas as escolas cadastradas', ''],
    ['Ações', 'Botões para editar, visualizar, ativar/desativar', '']
])

create_test_table(doc, '5.2 Dashboard Administrativo', [
    ['Elemento', 'Esperado', 'Status'],
    ['Acesso', 'Admin consegue acessar /admin-dashboard', ''],
    ['Resumo', 'Nome da escola, dados, ano letivo', ''],
    ['Seções', 'Turmas, Alunos, Professores, Configurações', ''],
    ['Métricas', 'Total de alunos, turmas, ausências', ''],
    ['Restrição', 'Admin não vê dados de outras escolas', '']
])

create_test_table(doc, '5.3 Dashboard Financeiro', [
    ['Elemento', 'Esperado', 'Status'],
    ['Acesso', 'Financeiro consegue acessar /fin-dashboard', ''],
    ['Resumo', 'Receita, boletos pendentes, pagos', ''],
    ['Filtros', 'Por mês, status, tipo de pagamento', ''],
    ['Relatórios', 'Exportação de dados em CSV/PDF', ''],
    ['Asaas', 'PIX codes gerados e exibidos corretamente', '']
])

# 6. Class Management
doc.add_heading('6. Testes de Gestão de Turmas', level=1)

create_test_table(doc, '6.1 Criar Nova Turma', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Acessar seção de Turmas', 'Lista de turmas é exibida'],
    ['2', 'Clicar em \'Adicionar Turma\'', 'Modal/formulário abre'],
    ['3', 'Preencher dados', 'Campos aceitam entrada'],
    ['4', 'Selecionar disciplinas', 'Checkboxes aparecem e são selecionáveis'],
    ['5', 'Clicar em \'Salvar\'', 'Turma é criada'],
    ['✓', 'Teste Passado', 'Nova turma aparece na lista']
])

create_test_table(doc, '6.2 Editar Turma', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Localizar uma turma', 'Turma está listada'],
    ['2', 'Clicar em \'Editar\'', 'Modal de edição abre com dados atuais'],
    ['3', 'Modificar dados', 'Campos permitem edição'],
    ['4', 'Clicar em \'Salvar\'', 'Turma é atualizada'],
    ['5', 'Verificar mudanças', 'Dados atualizados aparecem na lista'],
    ['✓', 'Teste Passado', 'Turma é editada corretamente']
])

create_test_table(doc, '6.3 Adicionar Alunos à Turma', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Selecionar uma turma', 'Detalhes são exibidos'],
    ['2', 'Clicar em \'Adicionar Alunos\'', 'Modal abre'],
    ['3', 'Selecionar alunos', 'Alunos são selecionáveis'],
    ['4', 'Clicar em \'Confirmar\'', 'Alunos são adicionados'],
    ['5', 'Verificar lista', 'Alunos aparecem na turma'],
    ['✓', 'Teste Passado', 'Alunos adicionados corretamente']
])

# 7. Student Management
doc.add_heading('7. Testes de Gestão de Alunos', level=1)

create_test_table(doc, '7.1 Criar Novo Aluno', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Acessar seção de Alunos', 'Lista de alunos é exibida'],
    ['2', 'Clicar em \'Adicionar Aluno\'', 'Formulário abre'],
    ['3', 'Preencher dados completos', 'Campos aceitam entrada'],
    ['4', 'Validar email/telefone', 'Email e telefone são validados'],
    ['5', 'Clicar em \'Salvar\'', 'Aluno é criado'],
    ['✓', 'Teste Passado', 'Novo aluno aparece na lista']
])

# 8. Financial Management
doc.add_heading('8. Testes de Gestão Financeira', level=1)

create_test_table(doc, '8.1 Gerar Boletos', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Acessar Dashboard Financeiro', 'Painel é exibido'],
    ['2', 'Selecionar período/turma', 'Opções de filtro aparecem'],
    ['3', 'Clicar em \'Gerar Boletos\'', 'Boletos são criados no Asaas'],
    ['4', 'Aguardar processamento', 'Mensagem de progresso aparece'],
    ['5', 'Verificar tabela', 'Boletos listados com status \'Pendente\''],
    ['✓', 'Teste Passado', 'Boletos gerados com sucesso']
])

create_test_table(doc, '8.2 Gerar PIX em Massa', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Filtrar boletos pendentes', 'Lista aparece'],
    ['2', 'Clicar em \'Enviar PIX em Massa\'', 'Confirmação aparece'],
    ['3', 'Confirmar ação', 'PIX são gerados'],
    ['4', 'Verificar progresso', 'Barra de progresso (X de Y) é exibida'],
    ['5', 'Verificar entrega', 'Mensagens/notificações foram entregues'],
    ['✓', 'Teste Passado', 'PIX em massa enviado com sucesso']
])

create_test_table(doc, '8.3 Enviar PIX Individual', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Localizar boleto pendente', 'Boleto está listado'],
    ['2', 'Clicar em \'PIX\'', 'QR code é gerado'],
    ['3', 'Verificar QR code', 'Code é legível e escaneável'],
    ['4', 'Clicar em \'Enviar via Mensagem\'', 'PIX é enviado'],
    ['5', 'Verificar recebimento', 'Mensagem com PIX aparece no chat'],
    ['✓', 'Teste Passado', 'PIX individual enviado com sucesso']
])

# 9. Communications
doc.add_heading('9. Testes de Comunicações', level=1)

create_test_table(doc, '9.1 Enviar Mensagem', [
    ['Passo', 'Ação', 'Resultado Esperado'],
    ['1', 'Acessar chat/mensagens', 'Interface é exibida'],
    ['2', 'Selecionar um pai/responsável', 'Conversa abre'],
    ['3', 'Digitar mensagem', 'Campo aceita entrada'],
    ['4', 'Clicar em \'Enviar\'', 'Mensagem é enviada'],
    ['5', 'Verificar histórico', 'Mensagem aparece no chat'],
    ['6', 'Verificar recebimento', 'Pai recebe a mensagem'],
    ['✓', 'Teste Passado', 'Mensagem enviada e recebida']
])

# 10. Responsiveness
doc.add_heading('10. Testes de Responsividade', level=1)

create_test_table(doc, '10.1 Desktop', [
    ['Elemento', 'Esperado', 'Status'],
    ['Layout', 'Bem distribuído em tela completa', ''],
    ['Sidebar', 'Visível e navegável', ''],
    ['Tabelas', 'Colunas bem dispostas, sem scroll horizontal', ''],
    ['Botões', 'Legíveis e clicáveis', '']
])

create_test_table(doc, '10.2 Mobile', [
    ['Elemento', 'Esperado', 'Status'],
    ['Layout', 'Stack vertical, otimizado para mobile', ''],
    ['Menu', 'Hamburger/drawer menu', ''],
    ['Formas', 'Campos verticais, fáceis de preencher', ''],
    ['Performance', 'Carregamento rápido em 3G/4G', '']
])

# 11. Security
doc.add_heading('11. Testes de Segurança', level=1)

create_test_table(doc, '11.1 Validação de Acesso', [
    ['Teste', 'Procedimento', 'Resultado Esperado'],
    ['RLS', 'Admin tenta acessar dados de outra escola', 'Acesso negado'],
    ['Token', 'Remover/alterar token no localStorage', 'Redirecionado para login'],
    ['CORS', 'Requisição de origem diferente', 'Requisição bloqueada']
])

create_test_table(doc, '11.2 Validação de Senhas', [
    ['Teste', 'Entrada', 'Resultado Esperado'],
    ['Fraca', '123456', 'Erro: não atende requisitos'],
    ['Válida', 'Senha@1234', 'Aceita e salva'],
    ['Hashing', 'Verificar armazenamento', 'Token como hash, não plain text']
])

# 12. Performance
doc.add_heading('12. Testes de Performance', level=1)
doc.add_paragraph('Tempo de carregamento do dashboard: < 2 segundos', style='List Bullet')
doc.add_paragraph('Tempo de resposta de requisição API: < 500ms', style='List Bullet')
doc.add_paragraph('Geração de boletos em massa: < 30 segundos para 50 boletos', style='List Bullet')
doc.add_paragraph('Sem erros de memória ou travamentos', style='List Bullet')

# 13. Final Checklist
doc.add_heading('13. Checklist Final de Lançamento', level=1)

checklist_items = [
    '✓ Todos os testes de autenticação passaram',
    '✓ Recuperação de senha funciona corretamente',
    '✓ Todos os dashboards carregam sem erros',
    '✓ CRUD de turmas funciona completamente',
    '✓ CRUD de alunos funciona completamente',
    '✓ Geração de boletos sem erros',
    '✓ PIX em massa funciona corretamente',
    '✓ Comunicações são entregues corretamente',
    '✓ Testes de responsividade aprovados',
    '✓ Testes de segurança aprovados',
    '✓ Performance dentro dos limites',
    '✓ Sem bugs críticos encontrados'
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
header_cells = table.rows[0].cells
header_cells[0].text = 'Item'
header_cells[1].text = 'Status'
shade_cell(header_cells[0], 'd5e8f0')
shade_cell(header_cells[1], 'd5e8f0')

for item in checklist_items:
    row_cells = table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = '☐'

# 14. Notes
doc.add_heading('14. Notas e Observações', level=1)
doc.add_paragraph('Use este espaço para anotações durante o teste:')
doc.add_paragraph('_' * 80)
doc.add_paragraph('_' * 80)
doc.add_paragraph('_' * 80)

# 15. Approval
doc.add_heading('15. Aprovação para Lançamento', level=1)

approval_table = doc.add_table(rows=3, cols=2)
approval_table.style = 'Table Grid'

approval_table.rows[0].cells[0].text = 'Testador'
approval_table.rows[0].cells[1].text = '_' * 40

approval_table.rows[1].cells[0].text = 'Data'
approval_table.rows[1].cells[1].text = '_' * 40

approval_table.rows[2].cells[0].text = 'Aprovado para Lançamento'
approval_table.rows[2].cells[1].text = '☐ SIM   ☐ NÃO'

doc.add_paragraph()
final = doc.add_paragraph('Documento de Teste Completado')
final.runs[0].font.bold = True
final.runs[0].font.italic = True

# Save document
doc.save('Plano_Testes_GestEscolar_MVP.docx')
print('✅ Documento de testes criado: Plano_Testes_GestEscolar_MVP.docx')
